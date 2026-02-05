"""
 ?????? ???? ?????
?????? ????? ? ?????? ????? ??????
"""

import os
import io
import tempfile
import asyncio
from typing import Optional, Dict, Any, List
from pathlib import Path
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

class FileManager:
    def __init__(self):
        self.supported_formats = {
            "text": [".txt", ".md", ".py", ".js", ".html", ".css", ".json"],
            "document": [".pdf", ".docx", ".doc"],
            "image": [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".gif"],
            "audio": [".wav", ".mp3", ".m4a", ".ogg", ".flac"]
        }
        
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        print(" ???? ???? ????????? ??")
    
    def get_file_type(self, filename: str) -> str:
        """????? ??? ????"""
        ext = Path(filename).suffix.lower()
        
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        
        return "unknown"
    
    def is_supported(self, filename: str) -> bool:
        """????? ???????? ????"""
        return self.get_file_type(filename) != "unknown"
    
    async def save_uploaded_file(self, file_data: bytes, filename: str) -> str:
        """????? ???? ????? ???"""
        try:
            # ????? ??? ????? ?? ???
            timestamp = int(asyncio.get_event_loop().time())
            safe_filename = f"{timestamp}_{filename}"
            file_path = self.upload_dir / safe_filename
            
            with open(file_path, "wb") as f:
                f.write(file_data)
            
            print(f" ???? ????? ??: {safe_filename}")
            return str(file_path)
            
        except Exception as e:
            print(f" ??? ?? ????? ????: {e}")
            raise
    
    async def read_text_file(self, file_path: str) -> str:
        """?????? ???? ????"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            print(f" ???? ???? ?????? ??: {len(content)} ???????")
            return content
            
        except UnicodeDecodeError:
            # ???? ?? encoding ??? ?????
            encodings = ["cp1256", "iso-8859-1", "windows-1252"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    print(f" ???? ?? encoding {encoding} ?????? ??")
                    return content
                except:
                    continue
            
            raise Exception("??????? ???? ?? ?? ??? encoding ?????")
        
        except Exception as e:
            print(f" ??? ?? ?????? ???? ????: {e}")
            raise
    
    async def read_pdf_file(self, file_path: str) -> str:
        """?????? ???? PDF"""
        try:
            content = ""
            
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
            
            print(f" PDF ?????? ??: {len(pdf_reader.pages)} ????")
            return content.strip()
            
        except Exception as e:
            print(f" ??? ?? ?????? PDF: {e}")
            raise
    
    async def read_docx_file(self, file_path: str) -> str:
        """?????? ???? Word"""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            print(f" Word document ?????? ??: {len(doc.paragraphs)} ????????")
            return content.strip()
            
        except Exception as e:
            print(f" ??? ?? ?????? Word: {e}")
            raise
    
    async def read_image_file(self, file_path: str) -> str:
        """?????? ??? ?? ????? (OCR)"""
        try:
            image = Image.open(file_path)
            
            # OCR ?? ???????? ?????
            text = pytesseract.image_to_string(
                image, 
                lang="fas+eng",  # ????? + ???????
                config="--psm 6"  # ????? ???? ??? ???????
            )
            
            print(f" OCR ????? ??: {len(text)} ???????")
            return text.strip()
            
        except Exception as e:
            print(f" ??? ?? OCR: {e}")
            # ??? tesseract ??? ?????
            if "tesseract" in str(e).lower():
                return "???: Tesseract OCR ??? ???? ???"
            raise
    
    async def process_file(self, file_path: str) -> Dict[str, Any]:
        """?????? ???? ????"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError("???? ???? ???")
            
            filename = os.path.basename(file_path)
            file_type = self.get_file_type(filename)
            file_size = os.path.getsize(file_path)
            
            result = {
                "filename": filename,
                "file_type": file_type,
                "file_size": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "supported": self.is_supported(filename),
                "content": "",
                "summary": "",
                "word_count": 0,
                "error": None
            }
            
            if not result["supported"]:
                result["error"] = "???? ???? ???????? ??????"
                return result
            
            # ?????? ????? ?? ???? ??? ????
            try:
                if file_type == "text":
                    content = await self.read_text_file(file_path)
                elif file_type == "document":
                    if filename.lower().endswith(".pdf"):
                        content = await self.read_pdf_file(file_path)
                    elif filename.lower().endswith((".docx", ".doc")):
                        content = await self.read_docx_file(file_path)
                    else:
                        content = ""
                elif file_type == "image":
                    content = await self.read_image_file(file_path)
                else:
                    content = ""
                
                result["content"] = content
                result["word_count"] = len(content.split()) if content else 0
                
                # ????????? ??? ????? ?????? ????
                if len(content) > 1000:
                    result["summary"] = await self.summarize_content(content)
                
            except Exception as e:
                result["error"] = str(e)
            
            return result
            
        except Exception as e:
            return {
                "filename": os.path.basename(file_path) if file_path else "??????",
                "error": str(e),
                "supported": False
            }
    
    async def summarize_content(self, content: str, max_length: int = 500) -> str:
        """????????? ?????"""
        try:
            # ????????? ???? ?? ???? ????? ???
            sentences = content.split(".")
            summary = ""
            
            for sentence in sentences[:5]:  # 5 ???? ???
                if len(summary + sentence) < max_length:
                    summary += sentence.strip() + ". "
                else:
                    break
            
            if len(summary) < 50:  # ??? ???? ????? ???
                summary = content[:max_length] + "..."
            
            return summary.strip()
            
        except:
            return content[:max_length] + "..." if len(content) > max_length else content
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """??????? ????"""
        try:
            if not os.path.exists(file_path):
                return {"error": "???? ???? ???"}
            
            stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            
            return {
                "filename": filename,
                "file_type": self.get_file_type(filename),
                "size": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "supported": self.is_supported(filename)
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    def cleanup_old_files(self, days: int = 7):
        """??? ???? ??????? ?????"""
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (days * 24 * 60 * 60)
            
            deleted_count = 0
            for file_path in self.upload_dir.glob("*"):
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    file_path.unlink()
                    deleted_count += 1
            
            print(f" {deleted_count} ???? ????? ??? ??")
            return deleted_count
            
        except Exception as e:
            print(f" ??? ?? ??? ???? ??????? ?????: {e}")
            return 0
    
    def get_status(self) -> Dict[str, Any]:
        """????? ???? ????"""
        try:
            upload_files = list(self.upload_dir.glob("*"))
            total_size = sum(f.stat().st_size for f in upload_files if f.is_file())
            
            return {
                "upload_dir": str(self.upload_dir),
                "supported_formats": self.supported_formats,
                "uploaded_files_count": len(upload_files),
                "total_size_mb": round(total_size / (1024 * 1024), 2)
            }
            
        except Exception as e:
            return {"error": str(e)}

# ????? ??????
file_manager = FileManager()
